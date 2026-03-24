import io
import json
import re
import requests
from flask import Flask, request, jsonify, Response, send_file, send_from_directory
from flask_cors import CORS

app = Flask(__name__)
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB

DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

SYSTEM_PROMPT = """你是一名具有二十年以上从业经验的资深专业编审，长期供职于高等教育出版机构，专注于经贸类、国际商务类教材的审稿与编校工作。你当前承担《外贸谈判策略与实战》教材的三审三校任务。该教材定位为高等院校经贸类专业核心课程教材，面向具备一定英语基础和商务背景的本科生或高职生，须严格符合正式出版教材的体例规范与学术标准。

本教材为中英双语教材，每个知识点均有中文表述和对应的英文表述。审校时须同步处理中英两个版本，确保两种语言版本在内容上互相对应、在语言风格上各自符合本语言的专业表达规范。

请你在不改动原文核心观点与知识框架的前提下，对提交的段落进行系统性语言审校与润色，使其达到可直接付印的出版水准。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
零、条目结构的彻底重写（绝对优先级，凌驾所有其他规则）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【触发条件】凡原文出现以下任意结构，不得在原有条目基础上逐条润色，必须彻底解构后重新拟写为流畅的学术散文段落：
・以 ①②③④⑤、（一）（二）（三）、第一/第二/第三 分条列举的内容块
・以 ·、•、-、— 开头的多行并列条目（连续 ≥ 3 条）
・短行一一对应列举（每行一个短语或短句，连续 ≥ 3 行）
・"包括以下几点：""主要体现在以下方面：""如下所示：" 之后接的条目格式
・英文中以 "1. 2. 3." 或 "First, Second, Third," 或 "- / •" 开头的并列项目

【三步重写法】
第一步·提取：从各条目中抽取实质内容，辨识条目间逻辑关系（并列/递进/因果/对比/条件）
第二步·融合：以该逻辑关系为骨架，将各条目内容融入 2～4 个完整散文段落，用句间语义衔接（"由此""在此基础上""相较而言""与之对应"等）替代序号
第三步·检验：确认输出文字中不含任何条目符号、数字序号、短行列举

【绝对禁止】
✗ 保留原条目格式，仅改写各条的措辞——这不是改写，是无效操作
✗ 将"①②③"改为"首先/其次/最后"——条目化逻辑未变，仍不可接受
✗ 在每条结尾加一个衔接句就算完成——条目结构依然存在
✗ 输出中出现 ①②③（一）（二）·•- 等任何条目符号

【示例对比】
❌ 错误（不可接受）——条目格式未解构，仅逐条润色：
原文：询盘的作用包括：①建立联系；②了解行情；③比较报价
改后：询盘的作用主要涵盖：①建立买卖双方的初步联系；②充分了解市场行情；③便于比较供应商报价

✅ 正确（应当输出）——彻底散文化，逻辑内化于句间：
询盘作为外贸交易的起始环节，其价值体现在关系、信息与决策三个层面：它建立了买卖双方的初步沟通渠道，使询盘方得以获取市场行情的基本参照；与此同时，多方报价的横向比较也为后续谈判提供了合理的价格预期。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
一、AIGC痕迹识别与消除（最高优先级）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【结构性套话 — 须删除或改写】
・删除"首先……其次……再次……最后……"式的枚举过渡，以自然的逻辑连接词或句间语义衔接替代
・删除段落结尾的总结套语："综上所述""总的来说""由此可见，……的重要性不言而喻"等
・不使用"值得注意的是""不难发现""需要指出的是""有必要强调"等作为段落跳板的虚置引导语
・避免"第一、第二、第三"与正文散文段落混用造成的结构割裂

【评价性空话 — 须删除】
・删除缺乏实质内容的定性语句："这一点至关重要""具有重要意义""发挥着不可忽视的作用"
・避免以"可以看出""我们不难理解"等伪结论句收束论述
・删除对显而易见内容的反复强调与冗余解释

【模板化句式 — 须改写】
・改写"A是B的基础，B是C的前提，C决定D"式链条堆叠
・避免同一段落内多句共用相同的句型框架（如"通过……可以……；通过……可以……"）
・改写"能够……，从而……，进而……"式流水线推进句，使逻辑关系更为内在、自然
・精简名词并列堆砌（如连列六七项的罗列），聚焦实质性核心表达

【英文AIGC标记 — 须消除】
・删除"It is worth noting that / It is important to emphasize that / Needless to say"等虚置引导
・删除"In summary, / To conclude, / Overall, / In a nutshell,"等程式化收尾
・改写"not only...but also / on the one hand...on the other hand"的机械对仗堆砌
・避免"plays a crucial/vital/pivotal role in / is of paramount importance"等固定搭配滥用
・消除被动语态与现在分词的过度连缀

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
二、学术性与严谨性
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【中文术语规范】
・统一使用本领域规范术语：询盘、报盘、还盘、交易条件、贸易惯例等
・专业术语首次出现可附英文对应词，格式：中文（英文）
・涉及国际贸易惯例（CISG、UCP 600、Incoterms等）须准确援引，不得含糊表述

【英文术语规范】
・使用国际商务领域通用术语：enquiry/inquiry, offer, counter-offer, trade terms等
・同一概念前后一致，不随意替换近义词（如enquiry与inquiry在全文中择一统一）
・避免过度使用被动语态堆叠，保持主动与被动的合理比例

【逻辑论证】
・每一知识点的陈述须具备内在逻辑自洽性，因果关系、层进关系须在句法层面体现
・各段落间的承接关系须自然流畅，避免"另外""同时"等万能衔接词的泛用

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
三、文体与语言规范
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【中文】
・全程使用正式书面语，禁止口语化表达（咱们、其实、就是说、说白了等）
・避免"非常""十分""极其"等程度副词的滥用
・教材正文不使用第一人称，采用客观陈述语气
・标点符号遵守国家标准 GB/T 15834

【英文】
・使用正式学术书面英语，将口语化缩写还原：don't→do not，it's→it is
・句子长度适中，避免过度复杂的嵌套从句
・保持第三人称客观视角，全文语体一致

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
四、输出格式要求（严格执行）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

・直接输出修改后的完整文段
・严格保持与原文相同的段落结构和顺序（中文段落在前，英文段落在后）
・每个自然段作为独立一行输出，段落之间用空行分隔
・不添加任何说明、注释、括号标注、"编审说明"或解释性文字
・不输出原文，只输出修改后的版本
・若某段无需修改，原文照录
・不添加任何Markdown格式符号（**、##、- 等）"""


# ─── helpers ────────────────────────────────────────────────────────────────

def is_cjk(c: str) -> bool:
    cp = ord(c)
    return (
        0x4E00 <= cp <= 0x9FFF
        or 0x3400 <= cp <= 0x4DBF
        or 0x20000 <= cp <= 0x2A6DF
        or 0xF900 <= cp <= 0xFAFF
    )


def detect_lang(text: str) -> str:
    if not text:
        return "en"
    cjk = sum(1 for c in text if is_cjk(c))
    return "zh" if cjk / max(len(text), 1) > 0.15 else "en"


def is_numbered_heading(text: str) -> bool:
    """Match patterns: 1.  /  1.1  /  1.2.3  (with trailing space + content)"""
    return bool(re.match(r"^\d+(\.\d+)*\.?\s+\S", text.strip()))


# ─── formatting helpers ──────────────────────────────────────────────────────

def _pt_or_none(length_obj):
    try:
        return round(float(length_obj.pt), 2) if length_obj is not None else None
    except Exception:
        return None


def _ea_font(run):
    """Return the East-Asia (CJK) font name from a run's XML, or None."""
    try:
        from docx.oxml.ns import qn
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return None
        rfonts = rpr.find(qn("w:rFonts"))
        return rfonts.get(qn("w:eastAsia")) if rfonts is not None else None
    except Exception:
        return None


def _fmt_from_para(para):
    """Extract the key formatting attributes from a docx paragraph."""
    # Character-level: read from the first non-empty run
    font_name = font_name_ea = font_size_pt = None
    bold = italic = None
    for run in para.runs:
        if not run.text:
            continue
        font_name = run.font.name
        font_name_ea = _ea_font(run)
        font_size_pt = _pt_or_none(run.font.size)
        bold = run.bold
        italic = run.italic
        break  # first run is enough

    # Paragraph-level
    pf = para.paragraph_format
    alignment = None
    try:
        if para.alignment is not None:
            alignment = para.alignment.value
    except Exception:
        pass

    return {
        "font_name":     font_name,
        "font_name_ea":  font_name_ea,
        "font_size_pt":  font_size_pt,
        "bold":          bold,
        "italic":        italic,
        "space_before":  _pt_or_none(pf.space_before),
        "space_after":   _pt_or_none(pf.space_after),
        "left_indent":   _pt_or_none(pf.left_indent),
        "first_line_indent": _pt_or_none(pf.first_line_indent),
        "alignment":     alignment,
    }


def parse_docx(file_obj):
    from docx import Document

    doc = Document(file_obj)
    groups: list[dict] = []
    current: dict | None = None
    gid = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_lower = para.style.name.lower()
        is_heading_style = "heading" in style_lower or "标题" in style_lower
        is_numbered = is_numbered_heading(text)

        fmt = _fmt_from_para(para)

        if is_heading_style or is_numbered:
            if current and current["paragraphs"]:
                groups.append(current)
            gid += 1
            current = {
                "id": gid,
                "heading": text,
                "paragraphs": [{"text": text, "lang": "heading",
                                 "style": para.style.name, "fmt": fmt}],
                "combined_text": text,
            }
        else:
            if current is None:
                gid += 1
                current = {
                    "id": gid,
                    "heading": "(前言)",
                    "paragraphs": [],
                    "combined_text": "",
                }
            lang = detect_lang(text)
            current["paragraphs"].append({"text": text, "lang": lang,
                                           "style": para.style.name, "fmt": fmt})
            sep = "\n\n" if current["combined_text"] else ""
            current["combined_text"] += sep + text

    if current and current["paragraphs"]:
        groups.append(current)

    return groups


# ─── routes ─────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    # Serve directly to bypass Jinja2 parsing (Vue uses {{ }} syntax too)
    return send_from_directory("templates", "index.html")


@app.route("/health")
def health():
    return jsonify({"status": "ok"})


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "未找到上传文件"}), 400

    f = request.files["file"]
    if not f.filename or not f.filename.endswith(".docx"):
        return jsonify({"error": "请上传 .docx 格式的文件"}), 400

    try:
        groups = parse_docx(f)
        if not groups:
            return jsonify({"error": "文档中未找到任何段落组，请确认文档包含编号标题（如 1. / 1.1）"}), 400
        return jsonify({"groups": groups, "total": len(groups)})
    except Exception as exc:
        return jsonify({"error": f"文件解析失败：{exc}"}), 500


@app.route("/review", methods=["POST"])
def review():
    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    temperature = float(data.get("temperature", 0.3))
    temperature = max(0.1, min(temperature, 1.0))
    api_key = request.headers.get("X-API-Key", "").strip()

    if not api_key:
        return jsonify({"error": "请先在右上角输入 API 密钥"}), 401
    if not text:
        return jsonify({"error": "文本内容为空"}), 400

    def generate():
        try:
            resp = requests.post(
                DEEPSEEK_API_URL,
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "deepseek-chat",
                    "messages": [
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": text},
                    ],
                    "temperature": temperature,
                    "stream": True,
                    "max_tokens": 4096,
                },
                stream=True,
                timeout=120,
            )

            if resp.status_code != 200:
                try:
                    err_body = resp.json()
                    msg = err_body.get("error", {}).get("message", f"HTTP {resp.status_code}")
                except Exception:
                    msg = f"HTTP {resp.status_code}"
                yield f"data: {json.dumps({'error': msg})}\n\n"
                return

            for raw in resp.iter_lines():
                if not raw:
                    continue
                line = raw.decode("utf-8") if isinstance(raw, bytes) else raw
                if not line.startswith("data: "):
                    continue
                payload = line[6:]
                if payload.strip() == "[DONE]":
                    yield "data: [DONE]\n\n"
                    return
                try:
                    chunk = json.loads(payload)
                    content = chunk["choices"][0]["delta"].get("content", "")
                    if content:
                        yield f"data: {json.dumps({'content': content})}\n\n"
                except Exception:
                    pass

        except requests.exceptions.Timeout:
            yield f"data: {json.dumps({'error': 'API 请求超时（120 s），请重试'})}\n\n"
        except Exception as exc:
            yield f"data: {json.dumps({'error': str(exc)})}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ─── export helpers ──────────────────────────────────────────────────────────

def _apply_pf(pfmt, fmt: dict):
    """Copy paragraph-format fields from fmt into a python-docx ParagraphFormat."""
    from docx.shared import Pt
    try:
        if fmt.get("space_before") is not None:
            pfmt.space_before = Pt(fmt["space_before"])
        if fmt.get("space_after") is not None:
            pfmt.space_after = Pt(fmt["space_after"])
        if fmt.get("left_indent") is not None:
            pfmt.left_indent = Pt(fmt["left_indent"])
        if fmt.get("first_line_indent") is not None:
            pfmt.first_line_indent = Pt(fmt["first_line_indent"])
        if fmt.get("alignment") is not None:
            pfmt.alignment = fmt["alignment"]
    except Exception:
        pass


def _set_font(run, fmt: dict):
    """Apply font name (ASCII + East-Asia) and size from fmt to a run."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if fmt.get("font_size_pt"):
        try:
            run.font.size = Pt(fmt["font_size_pt"])
        except Exception:
            pass

    ascii_f = fmt.get("font_name")
    ea_f = fmt.get("font_name_ea")
    if ascii_f or ea_f:
        try:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            if ascii_f:
                rfonts.set(qn("w:ascii"), ascii_f)
                rfonts.set(qn("w:hAnsi"), ascii_f)
            if ea_f:
                rfonts.set(qn("w:eastAsia"), ea_f)
        except Exception:
            pass


def _add_para(doc, style_name: str, fmt: dict, text: str,
              strike=False, color=None, force_italic=False):
    """Add a paragraph to doc with the given style, formatting, text, and character overrides."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style_name]
    except (KeyError, Exception):
        pass
    _apply_pf(p.paragraph_format, fmt)

    run = p.add_run(text)

    # Preserve original bold/italic (unless overridden below)
    if fmt.get("bold") is not None:
        run.bold = fmt["bold"]
    if not force_italic and fmt.get("italic") is not None:
        run.italic = fmt["italic"]

    _set_font(run, fmt)

    if strike:
        run.font.strike = True
    if color is not None:
        run.font.color.rgb = color
    if force_italic:
        run.italic = True

    return p


@app.route("/export", methods=["POST"])
def export_doc():
    from docx import Document
    from docx.shared import RGBColor

    data = request.get_json(silent=True) or {}
    results: list[dict] = data.get("results", [])
    mode: str = data.get("mode", "tracked")   # "tracked" | "clean"

    if not results:
        return jsonify({"error": "没有可导出的内容"}), 400

    doc = Document()
    # Remove the default empty paragraph added by python-docx
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    TEAL = RGBColor(0, 134, 134)
    BLACK = RGBColor(0, 0, 0)

    for idx, result in enumerate(results):
        originals: list[dict] = result.get("original", [])
        modified: str = result.get("modified", "")
        decision: str = result.get("decision", "approved")

        if mode == "clean":
            # ── Clean version: only final content, no strikethrough / colour markup ──
            if decision == "skipped" or not modified:
                for pd in originals:
                    _add_para(doc,
                              style_name=pd.get("style", "Normal"),
                              fmt=pd.get("fmt", {}),
                              text=pd["text"])
            else:
                # Map modified lines → original paragraph styles by positional index
                lines = [l for l in modified.split("\n") if l.strip()]
                for i, ln in enumerate(lines):
                    pd = originals[i] if i < len(originals) else (originals[-1] if originals else {})
                    _add_para(doc,
                              style_name=pd.get("style", "Normal"),
                              fmt=pd.get("fmt", {}),
                              text=ln)

        else:
            # ── Tracked version (default) ──
            if decision == "skipped":
                # Reproduce original paragraphs as-is, preserving all formatting
                for pd in originals:
                    _add_para(doc,
                              style_name=pd.get("style", "Normal"),
                              fmt=pd.get("fmt", {}),
                              text=pd["text"])
            else:
                # original paragraphs: same formatting + strikethrough + black
                for pd in originals:
                    _add_para(doc,
                              style_name=pd.get("style", "Normal"),
                              fmt=pd.get("fmt", {}),
                              text=pd["text"],
                              strike=True, color=BLACK)

                # modified text: same paragraph format as first body para, teal italic
                if modified:
                    base = next((pd for pd in originals if pd.get("lang") != "heading"),
                                originals[0] if originals else {})
                    base_style = base.get("style", "Normal")
                    base_fmt   = base.get("fmt", {})

                    for ln in (l for l in modified.split("\n") if l.strip()):
                        _add_para(doc,
                                  style_name=base_style,
                                  fmt=base_fmt,
                                  text=ln,
                                  color=TEAL, force_italic=True)

        # Blank separator paragraph between groups (not after the last one)
        if idx < len(results) - 1:
            sep = doc.add_paragraph()
            # Match the spacing of the last original paragraph for consistency
            if originals:
                last_fmt = originals[-1].get("fmt", {})
                _apply_pf(sep.paragraph_format, last_fmt)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name="revised.docx",
        mimetype=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    )


DETECT_PROMPT = """你是一名AIGC内容检测专家，专门识别中英双语学术教材中AI生成文本的痕迹。

任务：对提交的文本进行AIGC率检测，输出0-100的整数评分。
定义：0 = 完全人工写作，100 = 完全AI生成。

评分起点为10分（任何文本均存在的基础概率），在此基础上累加以下信号分值。

━━━━━━━━ 中文 · 极高权重信号 +18~25分/处 ━━━━━━━━

这类信号出现即为强烈AI痕迹，每处直接加分：

✦ 【过度分点】在本应是散文的段落内使用明显的条目结构：
  - 以①②③④⑤、（一）（二）（三）、第一/第二/第三 作为每条起头的平行列举
  - 以"·""•""-""—"开头的多行并列条目（≥3条）
  - 短行列举格式（每行一个要点，3行以上）代替完整论述段落
✦ 【链式枚举过渡】"首先……其次……再次……最后……"完整链条
✦ 【套路收尾句】"综上所述""总的来说""由此可见……重要性不言而喻""综上"单独成句
✦ 【虚置引导语】段落以此类短语开头："值得注意的是""不难发现""需要指出的是""有必要强调""不可忽视的是"
✦ 【结构性空话叠用】同段出现≥2处："至关重要""具有重要意义""发挥着不可忽视的作用""不可或缺"
✦ 【链条逻辑堆叠】"A是B的基础，B是C的前提，C决定D的走向"式三层以上链式推导
✦ 【过度对称并列】同一段落出现≥3组完全平行的"既……又……""一方面……另一方面……"结构

━━━━━━━━ 中文 · 高权重信号 +10~15分/处 ━━━━━━━━

✦ 同段相同句型框架重复≥3次（排比句组）
✦ 流水线推进句："能够……从而……进而……"
✦ 程度副词密集（非常/十分/极其/尤为/格外/相当，同段≥3次）
✦ 段尾程式化拔高："这充分说明了……的重要性""可见……意义之深远"
✦ 名词并列罗列≥5项且缺乏实质区分（如"包括A、B、C、D、E、F等六个方面"）

━━━━━━━━ 中文 · 中权重信号 +5~8分/处 ━━━━━━━━

✦ 万能衔接词泛用（另外/同时/此外/与此同时，同段≥3次）
✦ 口号式收尾（"……任重道远""……具有深远意义""……前景广阔"）
✦ 开头套路（"随着……的不断发展""在……背景下""……已成为……的重要课题"）
✦ 段落结构异常完整，每段都有引入→论述→收尾三段式，无逻辑留白

━━━━━━━━ 英文 · 极高权重信号 +18~25分/处 ━━━━━━━━

✦ 【过度分点】Bullet-point or numbered list structure used where flowing prose is expected:
  - Lines starting with "-", "•", "*", "1.", "2.", "(1)", "(2)" forming parallel items ≥3
  - "First, … Second, … Third, … Finally, …" enumeration chain
✦ "It is worth noting that / It should be noted that / It is important to emphasize / Needless to say"
✦ "In summary, / To conclude, / In conclusion, / Overall, / To sum up, / In a nutshell," as paragraph closer
✦ "plays a crucial/vital/pivotal/key role in / is of paramount/utmost importance / is indispensable"（同段≥2处）
✦ "not only…but also / on the one hand…on the other hand"（同段≥2处机械对仗）

━━━━━━━━ 英文 · 高权重信号 +10~15分/处 ━━━━━━━━

✦ 被动语态与现在分词过度连缀（同句超过3层）
✦ 相同句型框架重复（多句均以"This/These…will/can/should/enables"起头）
✦ 程度副词密集（significantly/substantially/considerably/undoubtedly/evidently，同段≥3次）
✦ "Furthermore, / Moreover, / Additionally, / In addition," 同段≥3次

━━━━━━━━ 英文 · 中权重信号 +5~8分/处 ━━━━━━━━

✦ "As mentioned above / As stated earlier / As discussed" 自我回指套话
✦ "This paper / This chapter / This section aims to / seeks to / endeavors to" 元叙述句
✦ 句末程式化升华（"…thus contributing to…", "…thereby promoting…"）

━━━━━━━━ 评分规则 ━━━━━━━━

1. 逐句逐段扫描，每发现一处信号即累加对应分值
2. 累计总分（含起点10分）超过100取100
3. 中英文均有内容时，分别评分后取加权平均（中文段落数：英文段落数）
4. 若文本极短（<40字）或仅为标题行，置信度设为 low，rate 取50
5. 仅评估语言风格特征，不评判内容本身的正确性
6. 宁可高判，不可漏判——有疑义的信号按高权重计算

━━━━━━━━ 输出格式（严格执行） ━━━━━━━━

直接输出以下JSON，不添加任何其他文字、markdown符号或解释：
{
  "rate": <0到100的整数>,
  "indicators": [<最显著的AIGC特征，最多5条，用中文简短描述具体位置和特征>],
  "confidence": "<high|medium|low>"
}"""


@app.route("/detect", methods=["POST"])
def detect():
    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    api_key = request.headers.get("X-API-Key-Detect", "").strip()

    if not api_key:
        return jsonify({"error": "请提供检测用 API 密钥"}), 401
    if not text:
        return jsonify({"rate": 0, "indicators": [], "confidence": "low"})

    try:
        resp = requests.post(
            DEEPSEEK_API_URL,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": DETECT_PROMPT},
                    {"role": "user", "content": text},
                ],
                "temperature": 0.1,
                "stream": False,
                "max_tokens": 512,
                "response_format": {"type": "json_object"},
            },
            timeout=30,
        )
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        result = json.loads(content)
        rate = max(0, min(100, int(result.get("rate", 50))))
        indicators = [str(x) for x in result.get("indicators", [])][:5]
        confidence = result.get("confidence", "medium")
        return jsonify({"rate": rate, "indicators": indicators, "confidence": confidence})
    except requests.exceptions.Timeout:
        return jsonify({"error": "检测请求超时（30s）"}), 504
    except (KeyError, ValueError, json.JSONDecodeError) as exc:
        # Model returned non-JSON or unexpected structure – return neutral score
        return jsonify({"rate": 50, "indicators": [], "confidence": "low",
                        "warning": f"结果解析失败：{exc}"})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


if __name__ == "__main__":
    app.run(debug=True, port=5000, threaded=True)
